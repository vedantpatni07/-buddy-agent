# Copyright 2025 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Tools for the Document Processor Sub-Agent."""

import base64
import io
import logging
from typing import Dict, Any
import pdfplumber
import docx
from google.adk.tools import ToolContext

logger = logging.getLogger(__name__)


async def extract_pdf_text(
    document_path: str,
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """Extract text from a PDF document."""
    try:
        # Load document artifact
        document_artifact = await tool_context.load_artifact(document_path)
        if not document_artifact or not document_artifact.inline_data:
            return {
                "success": False,
                "error": "Could not load PDF document artifact",
                "text": "",
                "metadata": {}
            }
        
        # Decode PDF content
        pdf_bytes = base64.b64decode(document_artifact.inline_data.data)
        
        with io.BytesIO(pdf_bytes) as pdf_file:
            text = ""
            metadata = {"pages": 0, "sections": []}
            
            with pdfplumber.open(pdf_file) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}\n"
                        metadata["sections"].append({
                            "page": page_num,
                            "characters": len(page_text),
                            "text_preview": page_text[:200] + "..." if len(page_text) > 200 else page_text
                        })
            
            logger.info(f"Successfully extracted text from PDF: {len(text)} characters, {metadata['pages']} pages")
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata,
                "document_type": "PDF",
                "characters_extracted": len(text)
            }
            
    except Exception as e:
        logger.error(f"Error extracting PDF text from {document_path}: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "text": "",
            "metadata": {},
            "document_type": "PDF"
        }


async def extract_word_text(
    document_path: str,
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """Extract text from a Word document."""
    try:
        # Load document artifact
        document_artifact = await tool_context.load_artifact(document_path)
        if not document_artifact or not document_artifact.inline_data:
            return {
                "success": False,
                "error": "Could not load Word document artifact",
                "text": "",
                "metadata": {}
            }
        
        # Decode Word document content
        doc_bytes = base64.b64decode(document_artifact.inline_data.data)
        
        with io.BytesIO(doc_bytes) as doc_file:
            doc = docx.Document(doc_file)
            text = ""
            metadata = {"paragraphs": 0, "sections": []}
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:
                    text += f"{para_text}\n"
                    metadata["sections"].append({
                        "paragraph": para_num,
                        "text": para_text[:200] + "..." if len(para_text) > 200 else para_text,
                        "characters": len(para_text)
                    })
            
            metadata["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
            
            logger.info(f"Successfully extracted text from Word document: {len(text)} characters, {metadata['paragraphs']} paragraphs")
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata,
                "document_type": "Word",
                "characters_extracted": len(text)
            }
            
    except Exception as e:
        logger.error(f"Error extracting Word text from {document_path}: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "text": "",
            "metadata": {},
            "document_type": "Word"
        }


async def validate_document(
    document_path: str,
    expected_type: str,
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """Validate document format and accessibility."""
    try:
        # Load document artifact
        document_artifact = await tool_context.load_artifact(document_path)
        if not document_artifact or not document_artifact.inline_data:
            return {
                "success": False,
                "error": "Could not load document artifact",
                "is_valid": False,
                "validation_details": {}
            }
        
        # Check document size
        document_bytes = base64.b64decode(document_artifact.inline_data.data)
        file_size = len(document_bytes)
        
        # Basic validation
        validation_details = {
            "file_size_bytes": file_size,
            "file_size_mb": round(file_size / (1024 * 1024), 2),
            "has_content": file_size > 0,
            "expected_type": expected_type,
            "mime_type": document_artifact.inline_data.mime_type if hasattr(document_artifact.inline_data, 'mime_type') else "unknown"
        }
        
        # Type-specific validation
        if expected_type.lower() == 'pdf':
            validation_details["is_pdf"] = document_bytes.startswith(b'%PDF')
        elif expected_type.lower() in ['word', 'docx', 'doc']:
            validation_details["is_word"] = document_bytes.startswith(b'PK')  # ZIP-based format
        
        is_valid = (
            validation_details["has_content"] and
            file_size > 100 and  # Minimum size check
            (validation_details.get("is_pdf", False) or validation_details.get("is_word", False))
        )
        
        logger.info(f"Document validation completed: {document_path} - Valid: {is_valid}")
        
        return {
            "success": True,
            "is_valid": is_valid,
            "validation_details": validation_details,
            "recommendations": [] if is_valid else ["Check document format and ensure it's not corrupted"]
        }
        
    except Exception as e:
        logger.error(f"Error validating document {document_path}: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "is_valid": False,
            "validation_details": {}
        }
