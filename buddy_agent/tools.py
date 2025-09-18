# Copyright 2025 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Tools for the Buddy Agent system.

This module contains tools for document processing, RAG retrieval, and Q&A generation.
"""

import base64
import io
import logging
from typing import Dict, List, Any, Optional
import pdfplumber
import docx
from google.adk.tools import ToolContext
from google.adk.tools.agent_tool import AgentTool

logger = logging.getLogger(__name__)


async def process_document(
    document_path: str,
    document_type: str,
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """
    Process a document (PDF or Word) and extract text content.
    
    Args:
        document_path: Path to the document file
        document_type: Type of document ('pdf' or 'word')
        tool_context: ADK tool context
    
    Returns:
        Dictionary containing extracted text and metadata
    """
    try:
        logger.info(f"Processing {document_type} document: {document_path}")
        
        # Load document artifact
        document_artifact = await tool_context.load_artifact(document_path)
        if not document_artifact or not document_artifact.inline_data:
            return {
                "success": False,
                "error": "Could not load document artifact",
                "text": "",
                "metadata": {}
            }
        
        # Decode document content
        document_bytes = base64.b64decode(document_artifact.inline_data.data)
        
        # Extract text based on document type
        if document_type.lower() == 'pdf':
            text, metadata = await _extract_pdf_text(document_bytes)
        elif document_type.lower() in ['word', 'docx', 'doc']:
            text, metadata = await _extract_word_text(document_bytes)
        else:
            return {
                "success": False,
                "error": f"Unsupported document type: {document_type}",
                "text": "",
                "metadata": {}
            }
        
        # Store processed document in session state
        processed_docs = tool_context.state.get("processed_documents", [])
        processed_docs.append({
            "path": document_path,
            "type": document_type,
            "text": text,
            "metadata": metadata
        })
        tool_context.state["processed_documents"] = processed_docs
        
        logger.info(f"Successfully processed document. Extracted {len(text)} characters.")
        
        return {
            "success": True,
            "text": text,
            "metadata": metadata,
            "characters_extracted": len(text),
            "document_path": document_path
        }
        
    except Exception as e:
        logger.error(f"Error processing document {document_path}: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "text": "",
            "metadata": {}
        }


async def _extract_pdf_text(document_bytes: bytes) -> tuple[str, Dict[str, Any]]:
    """Extract text from PDF document bytes."""
    try:
        with io.BytesIO(document_bytes) as pdf_file:
            text = ""
            metadata = {"pages": 0, "sections": []}
            
            with pdfplumber.open(pdf_file) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}\n"
                        metadata["sections"].append({
                            "page": page_num,
                            "characters": len(page_text)
                        })
            
            return text, metadata
            
    except Exception as e:
        logger.error(f"Error extracting PDF text: {str(e)}")
        return "", {"error": str(e)}


async def _extract_word_text(document_bytes: bytes) -> tuple[str, Dict[str, Any]]:
    """Extract text from Word document bytes."""
    try:
        with io.BytesIO(document_bytes) as doc_file:
            doc = docx.Document(doc_file)
            text = ""
            metadata = {"paragraphs": 0, "sections": []}
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:
                    text += f"{para_text}\n"
                    metadata["sections"].append({
                        "paragraph": para_num,
                        "text": para_text[:100] + "..." if len(para_text) > 100 else para_text
                    })
            
            metadata["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
            
            return text, metadata
            
    except Exception as e:
        logger.error(f"Error extracting Word text: {str(e)}")
        return "", {"error": str(e)}


async def create_document_corpus(
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """
    Create a searchable corpus from processed documents.
    
    Args:
        tool_context: ADK tool context
    
    Returns:
        Dictionary containing corpus creation status
    """
    try:
        processed_docs = tool_context.state.get("processed_documents", [])
        
        if not processed_docs:
            return {
                "success": False,
                "error": "No processed documents found. Please process documents first.",
                "corpus_id": None
            }
        
        # For now, create a simple text-based corpus
        # In production, this would integrate with Vertex AI RAG
        corpus_text = ""
        corpus_metadata = []
        
        for doc in processed_docs:
            corpus_text += f"\n\n=== {doc['path']} ===\n{doc['text']}\n"
            corpus_metadata.append({
                "path": doc['path'],
                "type": doc['type'],
                "characters": len(doc['text']),
                "metadata": doc['metadata']
            })
        
        # Store corpus in session state
        tool_context.state["document_corpus"] = {
            "text": corpus_text,
            "metadata": corpus_metadata,
            "total_documents": len(processed_docs),
            "total_characters": len(corpus_text)
        }
        tool_context.state["rag_retrieval_ready"] = True
        
        logger.info(f"Created corpus with {len(processed_docs)} documents")
        
        return {
            "success": True,
            "corpus_id": "local_corpus",
            "total_documents": len(processed_docs),
            "total_characters": len(corpus_text),
            "metadata": corpus_metadata
        }
        
    except Exception as e:
        logger.error(f"Error creating document corpus: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "corpus_id": None
        }


async def query_documents(
    question: str,
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """
    Query the document corpus for relevant content.
    
    Args:
        question: User's question
        tool_context: ADK tool context
    
    Returns:
        Dictionary containing relevant document sections
    """
    try:
        corpus = tool_context.state.get("document_corpus")
        
        if not corpus:
            return {
                "success": False,
                "error": "No document corpus found. Please process documents and create corpus first.",
                "relevant_sections": []
            }
        
        # Simple keyword-based search for now
        # In production, this would use semantic search with embeddings
        question_lower = question.lower()
        relevant_sections = []
        
        corpus_text = corpus["text"]
        corpus_metadata = corpus["metadata"]
        
        # Split corpus into sections and search for relevant content
        sections = corpus_text.split("===")
        
        for i, section in enumerate(sections[1:], 1):  # Skip first empty section
            if any(keyword in section.lower() for keyword in question_lower.split()):
                # Extract document name and content
                lines = section.strip().split('\n')
                doc_name = lines[0].strip() if lines else "Unknown Document"
                content = '\n'.join(lines[1:]).strip()
                
                if content:
                    relevant_sections.append({
                        "document": doc_name,
                        "content": content[:1000] + "..." if len(content) > 1000 else content,
                        "relevance_score": 0.8,  # Placeholder score
                        "section_index": i
                    })
        
        # Sort by relevance (in production, this would be done by the RAG system)
        relevant_sections.sort(key=lambda x: x["relevance_score"], reverse=True)
        
        logger.info(f"Found {len(relevant_sections)} relevant sections for question: {question}")
        
        return {
            "success": True,
            "question": question,
            "relevant_sections": relevant_sections[:5],  # Return top 5 most relevant
            "total_sections_found": len(relevant_sections)
        }
        
    except Exception as e:
        logger.error(f"Error querying documents: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "relevant_sections": []
        }


async def generate_answer(
    question: str,
    relevant_sections: List[Dict[str, Any]],
    tool_context: ToolContext = None,
) -> Dict[str, Any]:
    """
    Generate a comprehensive answer based on relevant document sections.
    
    Args:
        question: User's question
        relevant_sections: Relevant document sections from query_documents
        tool_context: ADK tool context
    
    Returns:
        Dictionary containing the generated answer and metadata
    """
    try:
        if not relevant_sections:
            return {
                "success": False,
                "answer": "I couldn't find relevant information in the processed documents to answer your question.",
                "sources": [],
                "confidence": "Low"
            }
        
        # Combine relevant sections for context
        context = "\n\n".join([
            f"**From {section['document']}:**\n{section['content']}"
            for section in relevant_sections
        ])
        
        # Generate answer using the context
        # In production, this would use a more sophisticated LLM call
        answer = f"""
Based on the processed documents, here's what I found regarding your question: "{question}"

{context}

Please note that this answer is based on the content available in the processed documents. If you need more specific information, please let me know and I can help you find additional details.
        """.strip()
        
        # Extract sources
        sources = [
            {
                "document": section["document"],
                "relevance_score": section["relevance_score"]
            }
            for section in relevant_sections
        ]
        
        # Determine confidence based on number and quality of relevant sections
        confidence = "High" if len(relevant_sections) >= 3 else "Medium" if len(relevant_sections) >= 1 else "Low"
        
        logger.info(f"Generated answer for question: {question}")
        
        return {
            "success": True,
            "answer": answer,
            "sources": sources,
            "confidence": confidence,
            "sections_used": len(relevant_sections)
        }
        
    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "answer": "I encountered an error while generating an answer. Please try again.",
            "sources": [],
            "confidence": "Low"
        }

